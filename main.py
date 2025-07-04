"""
DocxReplacer - 永久免费文档关键词替换工具
使用 tkinter 实现的无注册纯净版
"""

import os
import json
import re
import configparser
import platform
import traceback
import webbrowser
from datetime import datetime
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from docx import Document

class DocxReplacerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("DocxReplacer - 文档关键词替换工具")
        self.root.geometry("1100x800")
        self.root.configure(bg="#f0f0f0")

        # 应用信息
        self.app_info = {
            "name": "DocxReplacer",
            "version": "2.0.0",
            "author": "Easy",
            "website": "https://fdianshuo@gmail.com",
            "copyright": f"© {datetime.now().year} Jesus love us."
        }

        # 加载设置
        self.settings = self.load_settings()
        self.replacements = []
        self.current_ruleset = ""

        # 创建界面
        self.create_widgets()

        # 加载初始规则组
        self.load_initial_ruleset()

        # 设置窗口关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def create_widgets(self):
        # 创建主框架
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=BOTH, expand=True)

        # 标题区域
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=X, pady=(0, 20))

        ttk.Label(title_frame, text=f"{self.app_info['name']} v{self.app_info['version']}",
                 font=("Arial", 20, "bold")).pack(side=TOP)
        ttk.Label(title_frame, text=self.app_info["copyright"],
                 font=("Arial", 10)).pack(side=TOP)

        # 规则组管理区域
        ruleset_frame = ttk.LabelFrame(main_frame, text="规则组管理", padding=10)
        ruleset_frame.pack(fill=X, pady=10)

        ttk.Label(ruleset_frame, text="规则组:", font=("Arial", 12)).grid(row=0, column=0, sticky=W, padx=5)

        self.ruleset_combo = ttk.Combobox(ruleset_frame, font=("Arial", 12), width=30)
        self.ruleset_combo.grid(row=0, column=1, padx=5)
        self.ruleset_combo.bind("<<ComboboxSelected>>", self.on_ruleset_selected)

        ttk.Button(ruleset_frame, text="加载", command=self.load_ruleset, width=8).grid(row=0, column=2, padx=5)
        ttk.Button(ruleset_frame, text="保存", command=self.save_ruleset, width=8).grid(row=0, column=3, padx=5)
        ttk.Button(ruleset_frame, text="删除", command=self.delete_ruleset, width=8).grid(row=0, column=4, padx=5)

        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="文档处理", padding=10)
        file_frame.pack(fill=X, pady=10)

        ttk.Label(file_frame, text="输入文档:", font=("Arial", 12)).grid(row=0, column=0, sticky=W, padx=5, pady=5)
        self.input_entry = ttk.Entry(file_frame, font=("Arial", 12), width=60)
        self.input_entry.grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="浏览...", command=self.browse_input_file).grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(file_frame, text="输出文档:", font=("Arial", 12)).grid(row=1, column=0, sticky=W, padx=5, pady=5)
        self.output_entry = ttk.Entry(file_frame, font=("Arial", 12), width=60)
        self.output_entry.grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="浏览...", command=self.browse_output_file).grid(row=1, column=2, padx=5, pady=5)

        # 规则表区域
        rules_frame = ttk.LabelFrame(main_frame, text="替换规则", padding=10)
        rules_frame.pack(fill=BOTH, expand=True, pady=10)

        # 创建表格
        columns = ("original", "replacement")
        self.rules_table = ttk.Treeview(rules_frame, columns=columns, show="headings", height=8)

        # 设置列标题
        self.rules_table.heading("original", text="原文本")
        self.rules_table.heading("replacement", text="替换为")

        # 设置列宽
        self.rules_table.column("original", width=400, anchor=CENTER)
        self.rules_table.column("replacement", width=400, anchor=CENTER)

        # 添加滚动条
        scrollbar = ttk.Scrollbar(rules_frame, orient=VERTICAL, command=self.rules_table.yview)
        self.rules_table.configure(yscroll=scrollbar.set)

        # 布局表格和滚动条
        self.rules_table.grid(row=0, column=0, sticky=NSEW)
        scrollbar.grid(row=0, column=1, sticky=NS)

        # 配置网格布局
        rules_frame.grid_rowconfigure(0, weight=1)
        rules_frame.grid_columnconfigure(0, weight=1)

        # 规则操作按钮
        button_frame = ttk.Frame(rules_frame)
        button_frame.grid(row=1, column=0, columnspan=2, pady=10, sticky=EW)

        ttk.Button(button_frame, text="添加规则", command=self.add_rule).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="编辑规则", command=self.edit_rule).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="删除规则", command=self.delete_rule).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="清除规则", command=self.clear_rules).pack(side=LEFT, padx=5)

        # 执行按钮
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=X, pady=20)

        ttk.Button(action_frame, text="执行替换", command=self.run_replace,
                  style="Accent.TButton").pack(side=TOP, pady=10)

        # 输出区域
        output_frame = ttk.LabelFrame(main_frame, text="输出信息", padding=10)
        output_frame.pack(fill=BOTH, expand=True)

        self.output_text = Text(output_frame, wrap=WORD, font=("Consolas", 11), height=8)
        scrollbar = ttk.Scrollbar(output_frame, orient=VERTICAL, command=self.output_text.yview)
        self.output_text.configure(yscrollcommand=scrollbar.set)

        self.output_text.grid(row=0, column=0, sticky=NSEW)
        scrollbar.grid(row=0, column=1, sticky=NS)

        output_frame.grid_rowconfigure(0, weight=1)
        output_frame.grid_columnconfigure(0, weight=1)

        # 底部信息
        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(fill=X, pady=10)

        website_label = ttk.Label(footer_frame, text=self.app_info["website"],
                                 foreground="blue", cursor="hand2", font=("Arial", 10))
        website_label.pack(side=TOP)
        website_label.bind("<Button-1>", lambda e: webbrowser.open(self.app_info["website"]))

        ttk.Label(footer_frame, text=f"设置保存在: {self.get_app_data_dir()}",
                 font=("Arial", 8), foreground="gray").pack(side=TOP, pady=5)

        # 配置样式
        self.configure_styles()

    def configure_styles(self):
        style = ttk.Style()
        style.configure("TButton", font=("Arial", 12), padding=6)
        style.configure("Accent.TButton", font=("Arial", 12, "bold"), padding=8, foreground="white", background="#4b8bbe")
        style.map("Accent.TButton", background=[("active", "#3a6a9d")])
        style.configure("Treeview", font=("Arial", 12), rowheight=28)
        style.configure("Treeview.Heading", font=("Arial", 12, "bold"))

    def get_app_data_dir(self):
        """获取应用数据目录（跨平台）"""
        system = platform.system()

        if system == "Windows":
            app_data = os.getenv('APPDATA')
            app_dir = os.path.join(app_data, "DocxReplacer")
        elif system == "Darwin":  # macOS
            app_dir = os.path.expanduser("~/Library/Application Support/DocxReplacer")
        else:  # Linux
            app_dir = os.path.expanduser("~/.config/DocxReplacer")

        os.makedirs(app_dir, exist_ok=True)
        return app_dir

    def get_rules_dir(self):
        """获取规则组存储目录"""
        rules_dir = os.path.join(self.get_app_data_dir(), "rulesets")
        os.makedirs(rules_dir, exist_ok=True)
        return rules_dir

    def load_settings(self):
        """从配置文件加载设置"""
        config_file = os.path.join(self.get_app_data_dir(), "config.ini")
        settings = {
            'last_input_dir': '',
            'last_output_dir': '',
            'last_ruleset': ''
        }

        if not os.path.exists(config_file):
            return settings

        try:
            config = configparser.ConfigParser()
            config.read(config_file)

            if 'Settings' in config:
                settings['last_input_dir'] = config['Settings'].get('last_input_dir', '')
                settings['last_output_dir'] = config['Settings'].get('last_output_dir', '')
                settings['last_ruleset'] = config['Settings'].get('last_ruleset', '')
        except Exception as e:
            self.log(f"加载设置失败: {str(e)}")

        return settings

    def save_settings(self):
        """保存设置到配置文件"""
        config_file = os.path.join(self.get_app_data_dir(), "config.ini")

        config = configparser.ConfigParser()
        config['Settings'] = {
            'last_input_dir': self.settings['last_input_dir'],
            'last_output_dir': self.settings['last_output_dir'],
            'last_ruleset': self.current_ruleset
        }

        try:
            with open(config_file, 'w') as f:
                config.write(f)
        except Exception as e:
            self.log(f"保存设置失败: {str(e)}")

    def get_ruleset_list(self):
        """获取所有规则组列表"""
        rules_dir = self.get_rules_dir()
        rulesets = []

        for file in os.listdir(rules_dir):
            if file.endswith(".json"):
                ruleset_name = file[:-5]  # 去掉.json后缀
                rulesets.append(ruleset_name)

        return sorted(rulesets)

    def load_ruleset(self, ruleset_name):
        """从文件加载规则组"""
        rules_dir = self.get_rules_dir()
        ruleset_file = os.path.join(rules_dir, f"{ruleset_name}.json")

        if not os.path.exists(ruleset_file):
            return []

        try:
            with open(ruleset_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            self.log(f"加载规则组失败: {str(e)}")
            return []

    def save_ruleset(self, ruleset_name):
        """保存规则组到文件"""
        rules_dir = self.get_rules_dir()
        ruleset_file = os.path.join(rules_dir, f"{ruleset_name}.json")

        try:
            with open(ruleset_file, 'w', encoding='utf-8') as f:
                json.dump(self.replacements, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            self.log(f"保存规则组失败: {str(e)}")
            return False

    def delete_ruleset(self, ruleset_name):
        """删除规则组"""
        rules_dir = self.get_rules_dir()
        ruleset_file = os.path.join(rules_dir, f"{ruleset_name}.json")

        try:
            if os.path.exists(ruleset_file):
                os.remove(ruleset_file)
                return True
        except Exception as e:
            self.log(f"删除规则组失败: {str(e)}")

        return False

    def load_initial_ruleset(self):
        """加载初始规则组"""
        # 更新规则组列表
        ruleset_list = self.get_ruleset_list()
        self.ruleset_combo['values'] = ruleset_list

        if self.settings['last_ruleset'] and self.settings['last_ruleset'] in ruleset_list:
            self.current_ruleset = self.settings['last_ruleset']
            self.ruleset_combo.set(self.current_ruleset)
            self.replacements = self.load_ruleset(self.current_ruleset)
            self.update_rules_table()
            self.log(f"已加载规则组: {self.current_ruleset}")

    def update_rules_table(self):
        """更新规则表格"""
        # 清除现有数据
        for item in self.rules_table.get_children():
            self.rules_table.delete(item)

        # 添加新数据
        for i, rule in enumerate(self.replacements):
            self.rules_table.insert("", "end", iid=i, values=(rule[0], rule[1]))

    def browse_input_file(self):
        """浏览输入文件"""
        initial_dir = self.settings['last_input_dir'] or os.path.expanduser("~")
        file_path = filedialog.askopenfilename(
            title="选择输入文档",
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")],
            initialdir=initial_dir
        )

        if file_path:
            self.input_entry.delete(0, END)
            self.input_entry.insert(0, file_path)
            self.settings['last_input_dir'] = os.path.dirname(file_path)

    def browse_output_file(self):
        """浏览输出文件"""
        initial_dir = self.settings['last_output_dir'] or os.path.expanduser("~")
        file_path = filedialog.asksaveasfilename(
            title="保存输出文档",
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")],
            defaultextension=".docx",
            initialdir=initial_dir
        )

        if file_path:
            self.output_entry.delete(0, END)
            self.output_entry.insert(0, file_path)
            self.settings['last_output_dir'] = os.path.dirname(file_path)

    def add_rule(self):
        """添加新规则"""
        self.replacements.append(["", ""])
        self.update_rules_table()
        self.edit_rule(len(self.replacements) - 1)

    def edit_rule(self, index=None):
        """编辑规则"""
        if index is None:
            # 获取选中的规则
            selected = self.rules_table.selection()
            if not selected:
                messagebox.showinfo("提示", "请先选择一条规则")
                return
            index = int(selected[0])

        rule = self.replacements[index]

        # 创建编辑窗口
        edit_win = Toplevel(self.root)
        edit_win.title("编辑规则")
        edit_win.geometry("500x250")
        edit_win.transient(self.root)
        edit_win.grab_set()

        # 原文本
        ttk.Label(edit_win, text="原文本:", font=("Arial", 12)).pack(pady=(20, 5), padx=20, anchor=W)
        old_entry = ttk.Entry(edit_win, font=("Arial", 12), width=50)
        old_entry.pack(padx=20, fill=X)
        old_entry.insert(0, rule[0])

        # 替换文本
        ttk.Label(edit_win, text="替换为:", font=("Arial", 12)).pack(pady=(10, 5), padx=20, anchor=W)
        new_entry = ttk.Entry(edit_win, font=("Arial", 12), width=50)
        new_entry.pack(padx=20, fill=X)
        new_entry.insert(0, rule[1])

        # 按钮
        button_frame = ttk.Frame(edit_win)
        button_frame.pack(pady=20)

        def save_and_close():
            self.replacements[index] = [old_entry.get(), new_entry.get()]
            self.update_rules_table()
            edit_win.destroy()

        ttk.Button(button_frame, text="保存", command=save_and_close, width=10).pack(side=LEFT, padx=10)
        ttk.Button(button_frame, text="取消", command=edit_win.destroy, width=10).pack(side=LEFT, padx=10)

    def delete_rule(self):
        """删除选中的规则"""
        selected = self.rules_table.selection()
        if not selected:
            messagebox.showinfo("提示", "请先选择一条规则")
            return

        index = int(selected[0])
        self.replacements.pop(index)
        self.update_rules_table()

    def clear_rules(self):
        """清除所有规则"""
        if not self.replacements:
            return

        if messagebox.askyesno("确认", "确定要清除所有规则吗？"):
            self.replacements = []
            self.current_ruleset = ""
            self.ruleset_combo.set("")
            self.update_rules_table()
            self.log("已清除所有规则")

    def load_ruleset(self):
        """加载选中的规则组"""
        ruleset_name = self.ruleset_combo.get()
        if not ruleset_name:
            messagebox.showinfo("提示", "请先选择一个规则组")
            return

        new_rules = self.load_ruleset(ruleset_name)
        if new_rules:
            self.replacements = new_rules
            self.current_ruleset = ruleset_name
            self.update_rules_table()
            self.log(f"已加载规则组: {ruleset_name}")
            self.log(f"包含 {len(self.replacements)} 条替换规则")
        else:
            messagebox.showerror("错误", f"无法加载规则组: {ruleset_name}")

    def save_ruleset(self):
        """保存当前规则组"""
        ruleset_name = self.ruleset_combo.get()
        if not ruleset_name:
            # 如果没有选中规则组，弹出输入框
            ruleset_name = simpledialog.askstring("保存规则组", "请输入规则组名称:", parent=self.root)
            if not ruleset_name:
                return

        if self.save_ruleset(ruleset_name):
            self.current_ruleset = ruleset_name
            # 更新规则组列表
            ruleset_list = self.get_ruleset_list()
            self.ruleset_combo['values'] = ruleset_list
            self.ruleset_combo.set(ruleset_name)
            self.log(f"规则组 '{ruleset_name}' 保存成功!")
        else:
            messagebox.showerror("错误", f"保存规则组 '{ruleset_name}' 失败")

    def delete_ruleset(self):
        """删除选中的规则组"""
        ruleset_name = self.ruleset_combo.get()
        if not ruleset_name:
            messagebox.showinfo("提示", "请先选择一个规则组")
            return

        if messagebox.askyesno("确认删除", f'确定要删除规则组 "{ruleset_name}" 吗?'):
            if self.delete_ruleset(ruleset_name):
                # 更新规则组列表
                ruleset_list = self.get_ruleset_list()
                self.ruleset_combo['values'] = ruleset_list

                # 如果删除的是当前规则组，清除当前规则
                if ruleset_name == self.current_ruleset:
                    self.current_ruleset = ""
                    self.replacements = []
                    self.update_rules_table()

                self.ruleset_combo.set("")
                self.log(f"规则组 '{ruleset_name}' 已删除")
            else:
                messagebox.showerror("错误", f"删除规则组 '{ruleset_name}' 失败")

    def on_ruleset_selected(self, event):
        """当规则组选中时更新当前规则组名称"""
        self.current_ruleset = self.ruleset_combo.get()

    def run_replace(self):
        """执行替换操作"""
        input_file = self.input_entry.get()
        output_file = self.output_entry.get() or input_file

        if not input_file:
            messagebox.showinfo("提示", "请选择输入文件")
            return

        if not os.path.exists(input_file):
            messagebox.showerror("错误", f"文件不存在: {input_file}")
            return

        if not self.replacements:
            messagebox.showinfo("提示", "请添加替换规则")
            return

        # 转换为字典格式
        replacements_dict = {rule[0]: rule[1] for rule in self.replacements}

        # 显示处理信息
        self.log("=" * 80)
        self.log(f"开始处理: {input_file}")
        self.log(f"输出到: {output_file}")
        self.log(f"使用 {len(replacements_dict)} 条替换规则...")
        self.log("-" * 80)

        try:
            success, message = self.replace_in_docx(input_file, output_file, replacements_dict)

            if success:
                self.log("替换成功!")
                self.log(f"输出文件: {output_file}")
                self.log("=" * 80)
                messagebox.showinfo("完成", f"替换成功!\n输出文件: {output_file}")
            else:
                self.log(f"错误: {message}")
                self.log("=" * 80)
                messagebox.showerror("错误", f"处理失败: {message.splitlines()[0]}")
        except Exception as e:
            self.log(f"发生错误: {str(e)}")
            self.log(traceback.format_exc())
            messagebox.showerror("错误", f"处理失败: {str(e)}")

    def replace_in_docx(self, input_path, output_path, replacements):
        """
        执行文档关键词替换
        :param input_path: 输入文件路径
        :param output_path: 输出文件路径
        :param replacements: 替换规则字典 {旧词: 新词}
        """
        try:
            doc = Document(input_path)

            # 处理段落
            for para in doc.paragraphs:
                for old, new in replacements.items():
                    # 不区分大小写的普通替换
                    pattern = re.compile(re.escape(old), flags=re.IGNORECASE)
                    para.text = pattern.sub(new, para.text)

            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for old, new in replacements.items():
                                pattern = re.compile(re.escape(old), flags=re.IGNORECASE)
                                para.text = pattern.sub(new, para.text)

            # 处理页眉
            for section in doc.sections:
                for header in section.header.paragraphs:
                    for old, new in replacements.items():
                        pattern = re.compile(re.escape(old), flags=re.IGNORECASE)
                        header.text = pattern.sub(new, header.text)

            doc.save(output_path)
            return True, "替换成功"
        except Exception as e:
            return False, f"错误: {str(e)}\n\n{traceback.format_exc()}"

    def log(self, message):
        """添加日志到输出区域"""
        self.output_text.insert(END, message + "\n")
        self.output_text.see(END)

    def on_closing(self):
        """窗口关闭事件处理"""
        self.save_settings()
        self.root.destroy()

if __name__ == "__main__":
    root = Tk()
    app = DocxReplacerApp(root)
    root.mainloop()